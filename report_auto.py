from docx import Document
import json

def create_report(date="240526"):
    doc = Document()
    doc.add_heading(f"MMF Command Trend Report - {date}", 0)
    data = json.load(open(f"meta/analysis_commands_{date}.json", "r", encoding="utf-8"))
    doc.add_heading("TOP 5 Commands", level=1)
    for cmd, cnt in data["top_commands"]:
        doc.add_paragraph(f"{cmd}: {cnt}회")
    doc.add_heading("클러스터 분석", level=1)
    for cl, cmds in data["clusters"].items():
        doc.add_paragraph(f"{cl}: {', '.join(cmds)}")
    doc.save(f"MMF_Report_{date}.docx")
    print(f"리포트 생성 완료: MMF_Report_{date}.docx")

# pip install python-docx 필요
if __name__ == "__main__":
    create_report()
